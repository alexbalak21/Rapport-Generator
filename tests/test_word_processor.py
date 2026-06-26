import tempfile
import unittest
from pathlib import Path

from docx import Document

from app.core.word_processor import WordProcessor


class WordProcessorTests(unittest.TestCase):
    def test_fill_placeholders_replaces_braced_placeholders(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "output.docx"

            doc = Document()
            doc.add_paragraph("Inspection Date: {{date}}")
            doc.save(template_path)

            processor = WordProcessor(str(template_path))
            processor.fill_placeholders({"{{date}}": "2024-06-12"}, str(output_path))

            output_doc = Document(output_path)
            self.assertIn("2024-06-12", output_doc.paragraphs[0].text)
            self.assertNotIn("{{date}}", output_doc.paragraphs[0].text)


if __name__ == "__main__":
    unittest.main()
