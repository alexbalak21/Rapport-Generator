import re

try:
    from docx import Document
    from docx.oxml.ns import qn
except Exception:
    Document = None
    qn = None

PLACEHOLDER_RE = re.compile(r"\{\{.*?\}\}")


class WordProcessor:
    def __init__(self, filepath):
        if Document is None:
            raise ImportError("python-docx is required for WordProcessor")
        self.filepath = filepath
        self.document = Document(filepath)

    # ------------------------------------------------------------------
    # Iterate every paragraph in the document including tables
    # ------------------------------------------------------------------

    def _all_paragraphs(self):
        yield from self.document.paragraphs
        for table in self.document.tables:
            yield from self._paragraphs_in_table(table)

    def _paragraphs_in_table(self, table):
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested in cell.tables:
                    yield from self._paragraphs_in_table(nested)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def extract_placeholders(self) -> list:
        found = []
        for para in self._all_paragraphs():
            found.extend(PLACEHOLDER_RE.findall(para.text))
        return [placeholder.strip("{}") for placeholder in found]

    def fill_placeholders(self, mapping: dict, output_path: str) -> None:
        for para in self._all_paragraphs():
            self._replace_in_paragraph(para, mapping)
            if self._has_remaining_placeholder(para.text):
                self._replace_in_paragraph(para, mapping)
        self.document.save(output_path)

    # ------------------------------------------------------------------
    # Core replacement — preserves per-run formatting
    # ------------------------------------------------------------------

    def _replace_in_paragraph(self, paragraph, mapping: dict) -> None:
        if not paragraph.runs:
            return

        if not self._has_remaining_placeholder("".join(r.text for r in paragraph.runs)):
            return

        self._merge_split_placeholders(paragraph)
        self._replace_runs(paragraph.runs, mapping)

    def _replace_runs(self, runs, mapping: dict) -> None:
        for run in runs:
            if not self._has_remaining_placeholder(run.text):
                continue
            run.text = self._replace_run_text(run.text, mapping)

    def _replace_run_text(self, text: str, mapping: dict) -> str:
        for placeholder, value in mapping.items():
            if placeholder in text:
                text = text.replace(placeholder, str(value))
        return text

    @staticmethod
    def _has_remaining_placeholder(text: str) -> bool:
        return "{{" in text and "}}" in text

    def _merge_split_placeholders(self, paragraph) -> None:
        runs = paragraph.runs
        i = 0
        while i < len(runs):
            text = runs[i].text
            open_idx = text.find("{{")
            if open_idx == -1:
                i += 1
                continue

            if "}}" in text[open_idx:]:
                i += 1
                continue

            j = i + 1
            accumulated = text
            while j < len(runs):
                accumulated += runs[j].text
                if "}}" in accumulated:
                    break
                j += 1

            if "}}" not in accumulated:
                i += 1
                continue

            runs[i].text = accumulated
            for k in range(i + 1, j + 1):
                runs[k].text = ""
            i += 1
